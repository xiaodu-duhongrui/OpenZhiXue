"""
文档解析服务
提供PDF和Word文档的解析和文本提取功能
支持扫描件 OCR 识别
"""

import os
import logging
import asyncio
from pathlib import Path
from typing import Optional, List, Dict, Any
from io import BytesIO

from fastapi import HTTPException

logger = logging.getLogger(__name__)


class DocumentParseError(Exception):
    """文档解析错误"""
    pass


class DocumentService:
    """文档解析服务类"""

    def __init__(self):
        """初始化文档解析服务"""
        self._check_dependencies()

    def _check_dependencies(self) -> None:
        """检查必要的依赖库是否安装"""
        self.pypdf_available = False
        self.docx_available = False
        self.pil_available = False

        try:
            import PyPDF2
            self.pypdf_available = True
            logger.info("PyPDF2 is available")
        except ImportError:
            logger.warning("PyPDF2 is not installed. PDF parsing will not be available.")

        try:
            from docx import Document
            self.docx_available = True
            logger.info("python-docx is available")
        except ImportError:
            logger.warning("python-docx is not installed. Word document parsing will not be available.")

        try:
            from PIL import Image
            self.pil_available = True
            logger.info("Pillow is available")
        except ImportError:
            logger.warning("Pillow is not installed. Image parsing will not be available.")

    async def parse_document(self, file_path: str) -> Dict[str, Any]:
        """
        解析文档并提取文本内容

        Args:
            file_path: 文件路径

        Returns:
            包含解析结果的字典:
            - success: 是否成功
            - text: 提取的文本内容
            - page_count: 页数（PDF）
            - word_count: 字数
            - metadata: 文档元数据
            - error: 错误信息（如果失败）
            - is_scanned: 是否为扫描件（PDF）
            - ocr_used: 是否使用了 OCR
        """
        path = Path(file_path)
        if not path.exists():
            raise HTTPException(status_code=404, detail=f"文件不存在: {file_path}")

        extension = path.suffix.lower()

        if extension == ".pdf":
            return await self.parse_pdf(file_path)
        elif extension in [".docx", ".doc"]:
            return await self.parse_word(file_path)
        elif extension in [".png", ".jpg", ".jpeg"]:
            return await self.parse_image(file_path)
        else:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件类型: {extension}"
            )

    async def parse_pdf(self, file_path: str, use_ocr_for_scanned: bool = True) -> Dict[str, Any]:
        """
        解析PDF文档

        Args:
            file_path: PDF文件路径
            use_ocr_for_scanned: 如果检测到扫描件，是否使用 OCR

        Returns:
            包含解析结果的字典
        """
        if not self.pypdf_available:
            raise DocumentParseError("PyPDF2库未安装，无法解析PDF文件")

        try:
            import PyPDF2

            # 在线程池中执行同步IO操作
            result = await asyncio.to_thread(self._parse_pdf_sync, file_path, PyPDF2)
            
            # 检测是否为扫描件
            is_scanned = self._is_scanned_pdf(result)
            result["is_scanned"] = is_scanned
            result["ocr_used"] = False
            
            # 如果是扫描件且启用了 OCR，则使用 OCR 识别
            if is_scanned and use_ocr_for_scanned:
                logger.info(f"Detected scanned PDF, using OCR: {file_path}")
                try:
                    ocr_result = await self._ocr_pdf(file_path)
                    if ocr_result["success"]:
                        result["text"] = ocr_result["text"]
                        result["ocr_used"] = True
                        result["word_count"] = len(result["text"].replace(" ", "").replace("\n", ""))
                        logger.info(f"OCR completed for scanned PDF: {file_path}")
                except Exception as e:
                    logger.warning(f"OCR failed for scanned PDF: {e}, using original text")
            
            return result

        except Exception as e:
            logger.error(f"Failed to parse PDF {file_path}: {e}")
            raise DocumentParseError(f"PDF解析失败: {str(e)}")

    def _parse_pdf_sync(self, file_path: str, PyPDF2) -> Dict[str, Any]:
        """同步解析PDF文档"""
        text_content = []
        metadata = {}

        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)

            # 提取元数据
            if pdf_reader.metadata:
                metadata = {
                    "title": pdf_reader.metadata.get("/Title", ""),
                    "author": pdf_reader.metadata.get("/Author", ""),
                    "subject": pdf_reader.metadata.get("/Subject", ""),
                    "creator": pdf_reader.metadata.get("/Creator", ""),
                    "producer": pdf_reader.metadata.get("/Producer", ""),
                    "creation_date": str(pdf_reader.metadata.get("/CreationDate", "")),
                    "page_count": len(pdf_reader.pages),
                }

            # 提取每页文本
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- 第 {page_num} 页 ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {e}")
                    text_content.append(f"--- 第 {page_num} 页 ---\n[文本提取失败]")

        full_text = "\n\n".join(text_content)
        word_count = len(full_text.replace(" ", "").replace("\n", ""))

        return {
            "success": True,
            "text": full_text,
            "page_count": len(text_content),
            "word_count": word_count,
            "metadata": metadata,
            "error": None,
        }

    async def parse_word(self, file_path: str) -> Dict[str, Any]:
        """
        解析Word文档

        Args:
            file_path: Word文件路径

        Returns:
            包含解析结果的字典
        """
        if not self.docx_available:
            raise DocumentParseError("python-docx库未安装，无法解析Word文档")

        extension = Path(file_path).suffix.lower()

        if extension == ".doc":
            # .doc格式需要额外处理，python-docx只支持.docx
            raise DocumentParseError(
                ".doc格式暂不支持，请将文档转换为.docx格式后重试"
            )

        try:
            from docx import Document

            # 在线程池中执行同步IO操作
            result = await asyncio.to_thread(self._parse_docx_sync, file_path, Document)
            return result

        except Exception as e:
            logger.error(f"Failed to parse Word document {file_path}: {e}")
            raise DocumentParseError(f"Word文档解析失败: {str(e)}")

    def _parse_docx_sync(self, file_path: str, Document) -> Dict[str, Any]:
        """同步解析DOCX文档"""
        doc = Document(file_path)

        text_content = []
        metadata = {
            "core_properties": {},
            "paragraph_count": 0,
            "table_count": 0,
        }

        # 提取核心属性
        try:
            core_props = doc.core_properties
            metadata["core_properties"] = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or "",
            }
        except Exception as e:
            logger.warning(f"Failed to extract document properties: {e}")

        # 提取段落文本
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)

        metadata["paragraph_count"] = len(doc.paragraphs)
        metadata["table_count"] = len(doc.tables)

        # 提取表格文本
        if doc.tables:
            text_content.append("\n--- 表格内容 ---")
            for table_idx, table in enumerate(doc.tables, 1):
                text_content.append(f"\n表格 {table_idx}:")
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip(" |"):
                        text_content.append(row_text)

        full_text = "\n".join(text_content)
        word_count = len(full_text.replace(" ", "").replace("\n", ""))

        return {
            "success": True,
            "text": full_text,
            "page_count": 1,  # Word文档没有明确的页数概念
            "word_count": word_count,
            "metadata": metadata,
            "error": None,
        }

    async def extract_text_from_bytes(
        self,
        content: bytes,
        file_extension: str
    ) -> Dict[str, Any]:
        """
        从字节内容中提取文本

        Args:
            content: 文件内容字节
            file_extension: 文件扩展名（如 .pdf, .docx）

        Returns:
            包含解析结果的字典
        """
        extension = file_extension.lower()
        if not extension.startswith("."):
            extension = f".{extension}"

        if extension == ".pdf":
            return await self._extract_text_from_pdf_bytes(content)
        elif extension == ".docx":
            return await self._extract_text_from_docx_bytes(content)
        elif extension == ".doc":
            raise DocumentParseError(".doc格式暂不支持")
        else:
            raise DocumentParseError(f"不支持的文件类型: {extension}")

    async def _extract_text_from_pdf_bytes(self, content: bytes) -> Dict[str, Any]:
        """从PDF字节内容提取文本"""
        if not self.pypdf_available:
            raise DocumentParseError("PyPDF2库未安装")

        try:
            import PyPDF2

            result = await asyncio.to_thread(
                self._parse_pdf_bytes_sync,
                content,
                PyPDF2
            )
            return result

        except Exception as e:
            logger.error(f"Failed to parse PDF bytes: {e}")
            raise DocumentParseError(f"PDF解析失败: {str(e)}")

    def _parse_pdf_bytes_sync(self, content: bytes, PyPDF2) -> Dict[str, Any]:
        """同步从PDF字节解析"""
        text_content = []
        metadata = {}

        pdf_file = BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        if pdf_reader.metadata:
            metadata = {
                "title": pdf_reader.metadata.get("/Title", ""),
                "author": pdf_reader.metadata.get("/Author", ""),
                "page_count": len(pdf_reader.pages),
            }

        for page_num, page in enumerate(pdf_reader.pages, 1):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(f"--- 第 {page_num} 页 ---\n{page_text}")
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num}: {e}")

        full_text = "\n\n".join(text_content)
        word_count = len(full_text.replace(" ", "").replace("\n", ""))

        return {
            "success": True,
            "text": full_text,
            "page_count": len(text_content),
            "word_count": word_count,
            "metadata": metadata,
            "error": None,
        }

    async def _extract_text_from_docx_bytes(self, content: bytes) -> Dict[str, Any]:
        """从DOCX字节内容提取文本"""
        if not self.docx_available:
            raise DocumentParseError("python-docx库未安装")

        try:
            from docx import Document

            result = await asyncio.to_thread(
                self._parse_docx_bytes_sync,
                content,
                Document
            )
            return result

        except Exception as e:
            logger.error(f"Failed to parse DOCX bytes: {e}")
            raise DocumentParseError(f"Word文档解析失败: {str(e)}")

    def _parse_docx_bytes_sync(self, content: bytes, Document) -> Dict[str, Any]:
        """同步从DOCX字节解析"""
        docx_file = BytesIO(content)
        doc = Document(docx_file)

        text_content = []
        metadata = {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }

        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)

        if doc.tables:
            text_content.append("\n--- 表格内容 ---")
            for table_idx, table in enumerate(doc.tables, 1):
                text_content.append(f"\n表格 {table_idx}:")
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip(" |"):
                        text_content.append(row_text)

        full_text = "\n".join(text_content)
        word_count = len(full_text.replace(" ", "").replace("\n", ""))

        return {
            "success": True,
            "text": full_text,
            "page_count": 1,
            "word_count": word_count,
            "metadata": metadata,
            "error": None,
        }

    def get_supported_extensions(self) -> List[str]:
        """
        获取支持的文件扩展名列表

        Returns:
            支持的扩展名列表
        """
        extensions = []
        if self.pypdf_available:
            extensions.append(".pdf")
        if self.docx_available:
            extensions.extend([".docx", ".doc"])
        if self.pil_available:
            extensions.extend([".png", ".jpg", ".jpeg"])
        return extensions

    def is_supported(self, file_extension: str) -> bool:
        """
        检查文件类型是否支持

        Args:
            file_extension: 文件扩展名

        Returns:
            是否支持该文件类型
        """
        ext = file_extension.lower()
        if not ext.startswith("."):
            ext = f".{ext}"

        return ext in self.get_supported_extensions()

    def _is_scanned_pdf(self, parse_result: Dict[str, Any]) -> bool:
        """
        检测 PDF 是否为扫描件

        Args:
            parse_result: PDF 解析结果

        Returns:
            是否为扫描件
        """
        text = parse_result.get("text", "")
        word_count = parse_result.get("word_count", 0)
        page_count = parse_result.get("page_count", 1)

        # 如果平均每页字数少于 50 字，很可能是扫描件
        if page_count > 0:
            avg_words_per_page = word_count / page_count
            if avg_words_per_page < 50:
                logger.info(f"Low text density detected: {avg_words_per_page:.1f} words/page")
                return True

        # 如果文本内容非常少
        if word_count < 100:
            return True

        # 检查文本中是否包含大量乱码或不可读字符
        if text:
            # 计算可打印字符比例
            printable_chars = sum(1 for c in text if c.isprintable() or c in '\n\r\t')
            total_chars = len(text)
            if total_chars > 0:
                printable_ratio = printable_chars / total_chars
                if printable_ratio < 0.7:
                    logger.info(f"Low printable character ratio: {printable_ratio:.2f}")
                    return True

        return False

    async def _ocr_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        使用 OCR 识别 PDF 文档

        Args:
            file_path: PDF 文件路径

        Returns:
            OCR 结果
        """
        try:
            from app.services.ocr_service import get_ocr_service
            
            ocr_service = get_ocr_service()
            
            # 执行 OCR 识别
            results = await ocr_service.recognize_pdf(
                pdf_path=file_path,
                prompt="<image>\nExtract all text from this image and convert to markdown format.",
                enable_crop=True
            )

            # 合并所有页面的文本
            full_text = "\n\n".join(
                f"--- 第 {r.page_number} 页 ---\n{r.text}"
                for r in results
            )

            return {
                "success": True,
                "text": full_text,
                "page_count": len(results),
                "results": results
            }

        except Exception as e:
            logger.error(f"OCR failed for PDF {file_path}: {e}")
            return {
                "success": False,
                "text": "",
                "error": str(e)
            }

    async def parse_image(self, file_path: str) -> Dict[str, Any]:
        """
        解析图片文件（使用 OCR）

        Args:
            file_path: 图片文件路径

        Returns:
            包含解析结果的字典
        """
        if not self.pil_available:
            raise DocumentParseError("Pillow库未安装，无法解析图片文件")

        try:
            from app.services.ocr_service import get_ocr_service
            
            ocr_service = get_ocr_service()
            
            # 执行 OCR 识别
            result = await ocr_service.recognize_image_async(
                image_path=file_path,
                prompt="<image>\nExtract all text from this image and convert to markdown format.",
                enable_crop=True
            )

            return {
                "success": result.success,
                "text": result.text,
                "page_count": 1,
                "word_count": len(result.text.replace(" ", "").replace("\n", "")) if result.text else 0,
                "metadata": {
                    "file_type": "image",
                    "ocr_used": True,
                    "processing_time": result.processing_time
                },
                "error": result.error,
                "is_scanned": False,
                "ocr_used": True
            }

        except Exception as e:
            logger.error(f"Failed to parse image {file_path}: {e}")
            raise DocumentParseError(f"图片解析失败: {str(e)}")


# 创建全局文档服务实例
document_service = DocumentService()
